
from sentence_transformers import SentenceTransformer, util
import requests
from bs4 import BeautifulSoup
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
import os
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# --- Hàm lấy văn bản chính từ một URL (tìm tất cả thẻ <p>) ---
def fetch_text_from_url(url):
    try:
        response = requests.get(url)
        response.encoding = response.apparent_encoding
        if response.status_code != 200:
            return ""
        soup = BeautifulSoup(response.content, "html.parser")
        paragraphs = soup.find_all("p")
        text = "\n".join(p.get_text() for p in paragraphs)
        return text.strip()
    except Exception as e:
        print(f"Lỗi khi lấy URL: {e}")
        return ""

# --- Hàm chia văn bản thành các câu dựa trên dấu câu . ? ! ---
def split_into_sentences(text):
    import re
    sentences = re.split(r'(?<=[\.\?\!])\s+', text.strip())
    return [s for s in sentences if s]

# --- Khởi tạo mô hình SBERT ---
model = SentenceTransformer('all-MiniLM-L6-v2')

# --- Biến lưu kết quả và phần nghi ngờ ---
last_result = None
suspicious_parts = []

# --- Hàm kiểm tra đạo văn ---
def check_plagiarism():
    global last_result

    user_text = input_box.get("1.0", tk.END).strip()
    url = url_entry.get().strip()
    if not user_text or not url:
        result_label.config(text="Vui lòng nhập văn bản và URL nguồn.")
        return
    
    # Lấy nội dung từ URL
    source_text = fetch_text_from_url(url)
    if not source_text:
        result_label.config(text="Không thể tải văn bản từ URL.")
        return

    # Tính embedding toàn văn bản và cosine similarity
    embedding = model.encode([user_text, source_text])
    cos_sim = util.cos_sim(embedding[0], embedding[1])[0][0]
    similarity_percent = float(cos_sim) * 100
    result_label.config(text=f"Độ tương tự: {similarity_percent:.2f}%")

    # Lưu kết quả để xuất báo cáo
    last_result = {
        "url": url,
        "input_text": user_text,
        "similarity": round(similarity_percent, 2)
    }

    # Xóa highlight cũ
    input_box.tag_remove("highlight", "1.0", tk.END)
    
    # Tách câu và tính embedding từng câu để highlight
    user_sentences = split_into_sentences(user_text)
    source_sentences = split_into_sentences(source_text)
    user_emb = model.encode(user_sentences)
    source_emb = model.encode(source_sentences)
    threshold = 0.8  # Ngưỡng nghi ngờ
    suspicious_parts.clear()

    for i, sent in enumerate(user_sentences):
        if not sent.strip():
            continue
        cos_scores = util.cos_sim(user_emb[i], source_emb)[0]
        max_score = cos_scores.max().item()
        if max_score >= threshold:
            suspicious_parts.append(sent)
            pos = user_text.find(sent)
            if pos != -1:
                start_idx = f"1.0+{pos}c"
                end_idx = f"1.0+{pos+len(sent)}c"
                input_box.tag_add("highlight", start_idx, end_idx)
    
    # Cấu hình màu highlight
    input_box.tag_config("highlight", background="yellow")

# --- Xuất báo cáo PDF ---

# --- Xuất báo cáo PDF ---
def export_pdf():
    if not last_result:
        messagebox.showwarning("Chưa có dữ liệu", "Hãy kiểm tra đạo văn trước.")
        return

    # Try to register a TTF font that supports Unicode (Vietnamese)
    def _register_font():
        candidates = [
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/Library/Fonts/DejaVuSans.ttf",
            "/Library/Fonts/Arial Unicode.ttf",
            "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
            "/System/Library/Fonts/Supplemental/DejaVuSans.ttf",
            "/System/Library/Fonts/Arial.ttf",
            "/Library/Fonts/Times New Roman.ttf",
        ]
        for p in candidates:
            if os.path.exists(p):
                try:
                    pdfmetrics.registerFont(TTFont("AppFont", p))
                    return "AppFont"
                except Exception:
                    continue
        return None

    pdf = canvas.Canvas("KetQua_DaoVan.pdf", pagesize=A4)

    font_name = _register_font()

    # helper to ensure we don't pass unsupported unicode to builtin fonts
    def _safe_text(s, use_font):
        if use_font:
            return s
        try:
            return s.encode('latin-1', 'replace').decode('latin-1')
        except Exception:
            return s.encode('ascii', 'ignore').decode('ascii')
    
    if font_name:
        pdf.setFont(font_name, 12)
    else:
        # fallback to default (may not render Vietnamese); inform the user
        pdf.setFont("Helvetica", 12)
        messagebox.showwarning("Font", "Không tìm thấy font Unicode trên hệ thống. PDF có thể không hiển thị tiếng Việt đúng. Vui lòng cài DejaVuSans hoặc Noto Sans.")

    pdf.drawString(100, 800, _safe_text("BÁO CÁO KIỂM TRA ĐẠO VĂN", bool(font_name)))
    pdf.drawString(100, 770, _safe_text(f"URL gốc: {last_result['url']}", bool(font_name)))
    pdf.drawString(100, 740, _safe_text(f"Mức độ tương đồng: {last_result['similarity']}%", bool(font_name)))

    text_object = pdf.beginText(40, 700)
    if font_name:
        text_object.setFont(font_name, 12)
    else:
        text_object.setFont("Helvetica", 12)
    # split long text into lines for safety
    for line in ("Văn bản kiểm tra:",) + tuple(last_result['input_text'].splitlines()):
        text_object.textLine(_safe_text(line, bool(font_name)))
    pdf.drawText(text_object)
    pdf.save()

    messagebox.showinfo("Xuất PDF", "Đã lưu file KetQua_DaoVan.pdf thành công!")
def export_word():
    if not last_result:
        messagebox.showwarning("Chưa có dữ liệu", "Hãy kiểm tra đạo văn trước.")
        return

    doc = Document()
    doc.add_heading("BÁO CÁO KIỂM TRA ĐẠO VĂN", 0)
    doc.add_paragraph(f"URL gốc: {last_result['url']}")
    doc.add_paragraph(f"Mức độ tương đồng: {last_result['similarity']}%")
    doc.add_paragraph("Văn bản kiểm tra:")
    doc.add_paragraph(last_result['input_text'])
    doc.save("KetQua_DaoVan.docx")

    messagebox.showinfo("Xuất Word", "Đã lưu file KetQua_DaoVan.docx thành công!")

# --- Giao diện ---
root = tk.Tk()
root.title("🧠 AI Phát Hiện Đạo Văn")
root.geometry("700x650")

tk.Label(root, text="Nhập văn bản cần kiểm tra:", font=("Arial", 12, "bold")).pack(pady=5)
input_box = scrolledtext.ScrolledText(root, width=80, height=10, wrap=tk.WORD)
input_box.pack(pady=5)

tk.Label(root, text="Nhập URL nguồn để so sánh:", font=("Arial", 12, "bold")).pack(pady=5)
url_entry = tk.Entry(root, width=80)
url_entry.pack(pady=5)

tk.Button(root, text="🔎 Kiểm tra đạo văn", command=check_plagiarism, bg="#4CAF50", fg="white", font=("Arial", 12, "bold")).pack(pady=10)
result_label = tk.Label(root, text="", font=("Arial", 14, "bold"), fg="#d32f2f")
result_label.pack(pady=10)

tk.Button(root, text="📘 Xuất báo cáo PDF", command=export_pdf, bg="#2196F3", fg="white", font=("Arial", 12)).pack(pady=5)
tk.Button(root, text="📄 Xuất báo cáo Word", command=export_word, bg="#607D8B", fg="white", font=("Arial", 12)).pack(pady=5)

root.mainloop()
